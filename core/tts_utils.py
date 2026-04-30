"""
Text-to-Speech Utilities for Personal Drive
============================================
Handles text extraction from various file formats and audio generation
"""

import os
import io
import pyttsx3
from pathlib import Path
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

# Try to import optional dependencies
try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def extract_text_from_file(file_path, file_extension):
    """
    Extract text content from various file formats.
    
    Args:
        file_path: Path to the file or file-like object
        file_extension: File extension (.txt, .pdf, .docx, etc.)
    
    Returns:
        Extracted text string
    """
    extension = file_extension.lower()
    text = ""
    
    try:
        if extension == '.txt':
            # Handle text files
            if hasattr(file_path, 'read'):
                content = file_path.read()
                if isinstance(content, bytes):
                    text = content.decode('utf-8', errors='ignore')
                else:
                    text = content
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
        
        elif extension == '.pdf' and HAS_PDF:
            # Handle PDF files
            if hasattr(file_path, 'read'):
                pdf_reader = PdfReader(file_path)
            else:
                with open(file_path, 'rb') as f:
                    pdf_reader = PdfReader(f)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        elif extension == '.docx' and HAS_DOCX:
            # Handle Word documents
            if hasattr(file_path, 'read'):
                doc = DocxDocument(file_path)
            else:
                doc = DocxDocument(file_path)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
    
    except Exception as e:
        print(f"Error extracting text from {file_extension}: {str(e)}")
        return ""
    
    # Limit text to first 5000 characters to avoid very long TTS
    if len(text) > 5000:
        text = text[:5000] + "\n\n[טקסט מקוצר עקב אורך הקובץ]"
    
    return text.strip()


def generate_hebrew_audio_from_text(text):
    """
    Translate text to Hebrew (if needed), then generate Hebrew MP3 via gTTS.
    Requires internet access.

    Returns:
        bytes of MP3 audio, or None on failure
    """
    try:
        import io
        from gtts import gTTS
        from deep_translator import GoogleTranslator

        # Translate to Hebrew (auto-detect source language)
        print(f"[TTS-HE] Translating {len(text)} chars to Hebrew...")
        # deep-translator has a 5000-char limit per call; chunk if needed
        chunk_size = 4900
        if len(text) <= chunk_size:
            translated = GoogleTranslator(source='auto', target='iw').translate(text)
        else:
            chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
            translated = ' '.join(
                GoogleTranslator(source='auto', target='iw').translate(c) for c in chunks
            )
        print(f"[TTS-HE] Translation done ({len(translated)} chars). Generating audio...")

        tts = gTTS(text=translated, lang='iw', slow=False)  # Google uses 'iw' for Hebrew
        fp = io.BytesIO()
        tts.write_to_fp(fp)
        fp.seek(0)
        data = fp.read()
        print(f"[TTS-HE] Hebrew audio ready ({len(data)} bytes)")
        return data
    except Exception as e:
        import traceback
        print(f"[TTS-HE] Error generating Hebrew audio: {e}")
        traceback.print_exc()
        return None


def generate_audio_from_text(text, language='he', output_path=None):
    """
    Generate audio file from text using pyttsx3.
    
    Args:
        text: Text to convert to speech
        language: Language code (default: 'he' for Hebrew)
        output_path: Path to save audio file (if None, returns bytes)
    
    Returns:
        Path to audio file or bytes if output_path is None
    """
    if not text:
        return None
    
    try:
        import tempfile
        import platform
        
        print(f"[TTS] Initializing pyttsx3 engine...")
        engine = pyttsx3.init()
        
        # Set properties
        engine.setProperty('rate', 150)  # Speech rate
        engine.setProperty('volume', 0.9)  # Volume level (0.0 to 1.0)
        
        print(f"[TTS] Getting available voices...")
        # Try to find Hebrew voice, fallback to any available
        voices = engine.getProperty('voices')
        hebrew_voice = None
        
        print(f"[TTS] Found {len(voices)} voices available")
        for i, voice in enumerate(voices):
            print(f"  Voice {i}: {voice.name}")
            if 'he' in voice.languages or 'hebrew' in str(voice.name).lower():
                hebrew_voice = voice.id
                print(f"[TTS] Selected Hebrew voice: {voice.name}")
                break
        
        # Use first available voice if Hebrew not found
        if not hebrew_voice and voices:
            hebrew_voice = voices[0].id
            print(f"[TTS] Hebrew voice not found, using: {voices[0].name}")
        
        if hebrew_voice:
            engine.setProperty('voice', hebrew_voice)
        
        # Generate audio
        if output_path:
            print(f"[TTS] Saving audio to {output_path}")
            engine.save_to_file(text, output_path)
            engine.runAndWait()
            return output_path
        else:
            # Generate to temporary file, then read into memory (Windows-compatible)
            with tempfile.NamedTemporaryFile(suffix='.mp3', delete=False) as tmp:
                temp_path = tmp.name
            
            print(f"[TTS] Generating audio to temp file: {temp_path}")
            engine.save_to_file(text, temp_path)
            engine.runAndWait()
            
            # Read the generated file into memory
            if os.path.exists(temp_path):
                print(f"[TTS] Reading audio file...")
                with open(temp_path, 'rb') as f:
                    audio_data = f.read()
                
                # Clean up temp file
                try:
                    os.remove(temp_path)
                except:
                    pass
                
                print(f"[TTS] Audio generated successfully ({len(audio_data)} bytes)")
                return audio_data
            else:
                print(f"[TTS] ERROR: Temp file was not created at {temp_path}")
                return None
    
    except Exception as e:
        import traceback
        print(f"[TTS] ERROR generating audio: {str(e)}")
        traceback.print_exc()
        return None


def process_document_for_audio(document):
    """
    Process a document and generate audio file.
    Handles extracting text and generating MP3.
    
    Args:
        document: Document model instance
    
    Returns:
        Path to generated audio file or None
    """
    try:
        # Extract text from the document file
        text = extract_text_from_file(
            document.file,
            document.file_extension
        )
        
        if not text:
            return None
        
        # Generate audio file
        audio_filename = f"audio_{document.id}.mp3"
        audio_path = os.path.join('audio_files', audio_filename)
        
        # Create the audio
        audio_bytes = generate_audio_from_text(text, language='he')
        
        if audio_bytes:
            # Save to Django storage
            saved_path = default_storage.save(audio_path, ContentFile(audio_bytes))
            return saved_path
    
    except Exception as e:
        print(f"Error processing document for audio: {str(e)}")
        return None
